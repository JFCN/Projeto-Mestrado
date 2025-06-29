"""
Módulo 5: Sistema de Geração de Relatórios
Consolida resultados do pipeline em relatórios
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
import logging
import json
import matplotlib

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    docx_available = True
    print("✅ python-docx disponível para relatórios Word")
except ImportError:
    docx_available = False
    print("⚠️ python-docx não disponível. Relatórios Word desabilitados.")

logger = logging.getLogger('ComprehensiveReportGenerator')

class ComprehensiveReportGenerator:
    """Gerador de relatórios do sistema IA Tuberculose"""

    def __init__(self, results: Dict[str, Any], output_dir: Path, config: Dict[str, Any]):
        self.results = results
        self.output_dir = Path(output_dir)
        self.config = config
        self.timestamp = datetime.now()

        plt.style.use('default')
        sns.set_palette("husl")

        self.colors = {
            'primary': '#2E86AB',
            'secondary': '#A23B72',
            'success': '#F18F01',
            'warning': '#C73E1D',
            'text': '#2C3E50',
            'light': '#ECF0F1'
        }

        logger.info("📄 Inicializando gerador de relatórios...")
        logger.info(f"📁 Diretório: {self.output_dir}")
        logger.info(f"📊 Etapas disponíveis: {list(self.results.keys())}")

        self._diagnose_results()

    def _diagnose_results(self):
        """Diagnóstico dos dados disponíveis"""

        logger.info("🔍 DIAGNÓSTICO DOS RESULTADOS:")

        self.data_summary = {
            'total_steps': len(self.results),
            'available_steps': list(self.results.keys()),
            'step_details': {},
            'errors_found': [],
            'warnings': []
        }

        for step, data in self.results.items():
            step_info = {'type': type(data).__name__, 'content': {}}

            try:
                if isinstance(data, dict):
                    for key, value in data.items():
                        if key == 'timestamp':
                            step_info['timestamp'] = str(value)
                        elif isinstance(value, dict):
                            step_info['content'][key] = f"Dict com {len(value)} itens"
                        elif isinstance(value, list):
                            step_info['content'][key] = f"Lista com {len(value)} elementos"
                        elif hasattr(value, '__class__'):
                            step_info['content'][key] = value.__class__.__name__
                        else:
                            step_info['content'][key] = f"{type(value).__name__}"

                self.data_summary['step_details'][step] = step_info
                logger.info(f"   ✅ {step}: {len(step_info['content'])} elementos válidos")

            except Exception as e:
                error_msg = f"Erro ao analisar {step}: {str(e)}"
                self.data_summary['errors_found'].append(error_msg)
                logger.warning(f"   ⚠️ {error_msg}")

        if self.data_summary['errors_found']:
            logger.warning(f"🚨 {len(self.data_summary['errors_found'])} erros encontrados no diagnóstico")

        logger.info(f"✅ Diagnóstico concluído: {self.data_summary['total_steps']} etapas analisadas")

    def generate_word_report(self) -> Optional[str]:
        """Gera relatório em Word com formatação corporativa"""

        if not docx_available:
            logger.warning("⚠️ python-docx não disponível. Usando fallback.")
            return self._generate_text_report_fallback()

        try:
            logger.info("📄 Gerando relatório Word...")

            doc = Document()
            self._configure_professional_style(doc)
            self._add_cover_page(doc)
            self._add_executive_summary_professional(doc)
            self._add_pipeline_overview(doc)
            self._add_feature_analysis_detailed(doc)
            self._add_leakage_assessment_professional(doc)
            self._add_model_performance_detailed(doc)
            self._add_implementation_roadmap(doc)

            timestamp = self.timestamp.strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / "reports" / f"relatorio_completo_{timestamp}.docx"
            doc.save(str(filepath))

            logger.info(f"✅ Relatório Word gerado: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"❌ Erro ao gerar relatório Word: {e}")
            logger.error(f"   Tentando fallback para relatório em texto...")
            return self._generate_text_report_fallback()

    def _generate_text_report_fallback(self) -> str:
        """Fallback para relatório em texto quando Word falha"""

        try:
            timestamp = self.timestamp.strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / "reports" / f"relatorio_completo_{timestamp}.txt"

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("RELATÓRIO SISTEMA IA TUBERCULOSE INFANTIL\n")
                f.write("Sistema de Predição de Sobrevivência com IA\n")
                f.write("=" * 80 + "\n\n")

                self._write_fallback_content(f)

            logger.info(f"✅ Relatório fallback gerado: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"❌ Erro crítico no fallback: {e}")
            return ""

    def _configure_professional_style(self, doc):
        """Configura estilos para o documento"""

        try:
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

        except Exception as e:
            logger.warning(f"⚠️ Erro na configuração de estilo: {e}")

    def _add_cover_page(self, doc):
        """Página de capa"""

        try:
            title = doc.add_heading('RELATÓRIO DE ANÁLISE\nSISTEMA IA TUBERCULOSE INFANTIL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph('')

            subtitle = doc.add_heading('Sistema de Predição de Sobrevivência com Inteligência Artificial', 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph('')
            doc.add_paragraph('')

            info_para = doc.add_paragraph()
            info_para.add_run("Data de Execução: ").bold = True
            info_para.add_run(self.timestamp.strftime('%d/%m/%Y %H:%M:%S'))

            info_para = doc.add_paragraph()
            info_para.add_run("Pipeline Executado: ").bold = True
            info_para.add_run(f"{len(self.results)} de 5 etapas")

            info_para = doc.add_paragraph()
            info_para.add_run("Status Geral: ").bold = True
            info_para.add_run(self._get_overall_status())

            info_para = doc.add_paragraph()
            info_para.add_run("Risco Detectado: ").bold = True
            info_para.add_run(self._get_overall_risk())

            info_para = doc.add_paragraph()
            info_para.add_run("Recomendação: ").bold = True
            info_para.add_run(self._get_final_recommendation())

            if 'feature_selection' in self.results:
                fs_data = self.results['feature_selection']
                scenarios = fs_data.get('scenarios_processed', [])
                if scenarios:
                    info_para = doc.add_paragraph()
                    info_para.add_run("Cenários Analisados: ").bold = True
                    info_para.add_run(f"{len(scenarios)}")

                    info_para = doc.add_paragraph()
                    info_para.add_run("Cenários: ").bold = True
                    info_para.add_run(', '.join(scenarios))

            if 'model_training' in self.results:
                mt_data = self.results['model_training']
                if 'scenarios' in mt_data:
                    info_para = doc.add_paragraph()
                    info_para.add_run("Modelos Treinados: ").bold = True
                    info_para.add_run(f"{len(mt_data['scenarios'])}")

                if 'best_scenario' in mt_data and mt_data['best_scenario']:
                    info_para = doc.add_paragraph()
                    info_para.add_run("Melhor Modelo: ").bold = True
                    info_para.add_run(mt_data['best_scenario'])

            doc.add_page_break()

        except Exception as e:
            logger.warning(f"⚠️ Erro na página de capa: {e}")
            doc.add_heading('RELATÓRIO SISTEMA IA TUBERCULOSE', 0)
            doc.add_paragraph(f"Gerado em: {self.timestamp.strftime('%d/%m/%Y %H:%M:%S')}")

    def _add_executive_summary_professional(self, doc):
        """Resumo executivo"""

        try:
            doc.add_heading('RESUMO EXECUTIVO', 1)

            context_para = doc.add_paragraph()
            context_para.add_run("CONTEXTO: ").bold = True
            context_para.add_run("Desenvolvimento e validação de sistema de inteligência artificial para predição de sobrevivência em casos de tuberculose infantil, utilizando aprendizado de máquina com validação anti-vazamento temporal.")

            doc.add_paragraph('')

            status_para = doc.add_paragraph()
            status_para.add_run("STATUS DA EXECUÇÃO: ").bold = True
            completed_steps = len(self.results)
            status_para.add_run(f"{completed_steps}/5 etapas do pipeline concluídas com sucesso.")

            if 'feature_selection' in self.results:
                fs_data = self.results['feature_selection']
                scenarios = fs_data.get('scenarios_processed', [])

                if scenarios:
                    doc.add_paragraph('')
                    scenarios_para = doc.add_paragraph()
                    scenarios_para.add_run("CENÁRIOS ANALISADOS: ").bold = True
                    scenarios_para.add_run(f"{len(scenarios)}")

                    for scenario in scenarios:
                        doc.add_paragraph(f"• {scenario}")

            if 'model_training' in self.results:
                self._add_performance_summary(doc)

            doc.add_paragraph('')
            risk_heading = doc.add_paragraph()
            risk_heading.add_run("VALIDAÇÃO DE SEGURANÇA: ").bold = True

            risk_level = self._get_overall_risk()
            risk_para = doc.add_paragraph()
            risk_para.add_run("• Nível de Risco de Vazamento Temporal: ").bold = True
            risk_para.add_run(risk_level)

            risk_interpretation = self._get_leakage_interpretation(risk_level)

            doc.add_paragraph('')
            recommendation = self._get_final_recommendation()
            rec_para = doc.add_paragraph()
            rec_para.add_run("RECOMENDAÇÃO FINAL: ").bold = True
            rec_para.add_run(recommendation).bold = True

        except Exception as e:
            logger.warning(f"⚠️ Erro no resumo executivo: {e}")
            doc.add_heading('RESUMO EXECUTIVO', 1)
            doc.add_paragraph(f"Sistema executado com {len(self.results)} etapas concluídas.")
            doc.add_paragraph(f"Risco detectado: {self._get_overall_risk()}")
            doc.add_paragraph(f"Recomendação: {self._get_final_recommendation()}")

    def _add_performance_summary(self, doc):
        """Adiciona resumo de performance ao executivo"""
        try:
            mt_data = self.results['model_training']

            perf_para = doc.add_paragraph()
            perf_para.add_run("• Performance dos Modelos IA: ").bold = True

            best_scenario = None
            best_metrics = None
            models_count = 0

            if 'scenarios' in mt_data:
                models_count = len(mt_data['scenarios'])
                best_acc = 0

                for scenario, scenario_data in mt_data['scenarios'].items():
                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                        results = scenario_data['predictor'].results
                        acc = results.get('balanced_accuracy', 0)
                        if acc > best_acc:
                            best_acc = acc
                            best_scenario = scenario
                            best_metrics = results

            elif 'results' in mt_data:
                models_count = 1
                best_scenario = 'Principal'
                best_metrics = mt_data['results']

            if models_count > 0:
                models_para = doc.add_paragraph(f"  - {models_count} modelo(s) treinado(s) com sucesso")

            if best_metrics:
                acc = best_metrics.get('balanced_accuracy', 0)
                f1 = best_metrics.get('f1_score', 0)
                auc = best_metrics.get('auc_score', 0)

                best_para = doc.add_paragraph(f"  - Melhor Modelo ({best_scenario}): Acurácia {acc:.1%}, F1-Score {f1:.3f}, AUC {auc:.3f}")

                # Interpretação clínica
                if acc >= 0.80:
                    interp = "Excelente performance clínica"
                elif acc >= 0.70:
                    interp = "Boa performance clínica"
                elif acc >= 0.60:
                    interp = "Performance moderada"
                else:
                    interp = "Performance abaixo do esperado"

                doc.add_paragraph(f"  - Avaliação Clínica: {interp}")

        except Exception as e:
            logger.warning(f"⚠️ Erro no resumo de performance: {e}")

    def _add_pipeline_overview(self, doc):
        """Visão geral do pipeline executado"""

        try:
            doc.add_heading('STATUS DO PIPELINE', 1)

            completed_steps = len(self.results)
            doc.add_paragraph(f"PROGRESSO: {completed_steps}/5 etapas concluídas")
            doc.add_paragraph('')

            pipeline_steps = [
                ('feature_selection', '1. Seleção de Features'),
                ('clinical_analysis', '2. Análise Temporal'),
                ('leakage_detection', '3. Detecção de Data Leakage'),
                ('model_training', '4. Treinamento do Modelo'),
                ('report_generation', '5. Geração de Relatórios')
            ]

            for step_key, step_name in pipeline_steps:
                if step_key in self.results:
                    step_para = doc.add_paragraph(f"✅ {step_name}: CONCLUÍDA")

                    if step_key == 'feature_selection' and 'feature_selection' in self.results:
                        fs_data = self.results['feature_selection']
                        scenarios = fs_data.get('scenarios_processed', [])
                        doc.add_paragraph(f"   📊 Cenários processados: {len(scenarios)}")

                    elif step_key == 'model_training' and 'model_training' in self.results:
                        mt_data = self.results['model_training']
                        if 'scenarios' in mt_data:
                            models_count = len(mt_data['scenarios'])
                            doc.add_paragraph(f"   🤖 Modelos treinados: {models_count}")

                    elif step_key == 'leakage_detection' and 'leakage_detection' in self.results:
                        ld_data = self.results['leakage_detection']
                        risk = ld_data.get('overall_risk_level', 'UNKNOWN')
                        doc.add_paragraph(f"   🚨 Risco detectado: {risk}")
                else:
                    doc.add_paragraph(f"❌ {step_name}: NÃO EXECUTADA")

        except Exception as e:
            logger.warning(f"⚠️ Erro na visão geral: {e}")

    def _add_feature_analysis_detailed(self, doc):
        """Análise detalhada de features"""

        try:
            doc.add_heading('ANÁLISE DA SELEÇÃO DE FEATURES', 1)

            if 'feature_selection' not in self.results:
                doc.add_paragraph("Seleção de features não foi executada.")
                return

            fs_data = self.results['feature_selection']

            scenarios = fs_data.get('scenarios_processed', [])
            derived_features = fs_data.get('derived_features', [])

            summary_para = doc.add_paragraph()
            summary_para.add_run("CENÁRIOS PROCESSADOS: ").bold = True
            summary_para.add_run(f"{len(scenarios)}")

            summary_para = doc.add_paragraph()
            summary_para.add_run("FEATURES DERIVADAS CRIADAS: ").bold = True
            summary_para.add_run(f"{len(derived_features)}")

            if derived_features:
                doc.add_paragraph("Features derivadas:")
                for feature in derived_features:
                    doc.add_paragraph(f"• {feature}")

            if scenarios and hasattr(fs_data.get('selector'), 'get_data_for_scenario'):
                doc.add_paragraph('')
                scenarios_heading = doc.add_paragraph()
                scenarios_heading.add_run("Detalhamento por Cenário: ").bold = True

                for scenario in scenarios:
                    try:
                        data, features = fs_data['selector'].get_data_for_scenario(scenario)

                        doc.add_paragraph('')
                        scenario_para = doc.add_paragraph()
                        scenario_para.add_run(f"{scenario}: ").bold = True

                        doc.add_paragraph(f"• Registros: {len(data)}")
                        doc.add_paragraph(f"• Features selecionadas: {len(features)}")

                        if 'SITUA_ENCE' in data.columns:
                            target_dist = data['SITUA_ENCE'].value_counts()
                            doc.add_paragraph(f"• Distribuição target: {dict(target_dist)}")

                    except Exception as e:
                        doc.add_paragraph(f"{scenario}: Erro ao obter detalhes")
                        logger.warning(f"Erro ao obter dados do cenário {scenario}: {e}")

        except Exception as e:
            logger.warning(f"⚠️ Erro na análise de features: {e}")

    def _add_leakage_assessment_professional(self, doc):
        """Avaliação profissional de vazamento de dados"""

        try:
            doc.add_heading('ANÁLISE ANTI-LEAKAGE', 1)

            if 'leakage_detection' not in self.results:
                doc.add_paragraph("Detecção de vazamento não foi executada.")
                return

            ld_data = self.results['leakage_detection']

            overall_risk = ld_data.get('overall_risk_level', 'UNKNOWN')
            scenarios_analyzed = ld_data.get('scenarios_analyzed', [])

            assessment_para = doc.add_paragraph()
            assessment_para.add_run("RISCO GERAL: ").bold = True
            assessment_para.add_run(overall_risk)

            doc.add_paragraph(f"CENÁRIOS ANALISADOS: {len(scenarios_analyzed)}")

            if 'results_by_scenario' in ld_data:
                doc.add_paragraph('')
                detail_heading = doc.add_paragraph()
                detail_heading.add_run("Análise por Cenário: ").bold = True

                for scenario, results in ld_data['results_by_scenario'].items():
                    risk_level = results.get('risk_level', 'UNKNOWN')
                    suspicious_count = results.get('suspicious_count', 0)
                    total_features = results.get('total_features', 0)

                    doc.add_paragraph('')
                    scenario_para = doc.add_paragraph()
                    scenario_para.add_run(f"{scenario}: ").bold = True

                    doc.add_paragraph(f"• Risco: {risk_level}")
                    doc.add_paragraph(f"• Features suspeitas: {suspicious_count}/{total_features}")

                    if risk_level == 'LOW':
                        status = "✅ BAIXO RISCO"
                    elif risk_level == 'MODERATE':
                        status = "⚠️ RISCO MODERADO"
                    elif risk_level == 'HIGH':
                        status = "⚠️ ALTO RISCO"
                    else:
                        status = "❌ RISCO CRÍTICO"

                    doc.add_paragraph(f"• Status: {status}")

            interpretation = self._get_leakage_interpretation(overall_risk)
            doc.add_paragraph('')
            interp_para = doc.add_paragraph()
            interp_para.add_run("INTERPRETAÇÃO: ").bold = True
            doc.add_paragraph(interpretation)

        except Exception as e:
            logger.warning(f"⚠️ Erro na avaliação de vazamento: {e}")

    def _add_model_performance_detailed(self, doc):
        """Performance detalhada dos modelos"""

        try:
            doc.add_heading('PERFORMANCE DOS MODELOS', 1)

            if 'model_training' not in self.results:
                doc.add_paragraph("Treinamento de modelos não foi executado.")
                return

            mt_data = self.results['model_training']

            if 'scenarios' in mt_data:
                scenarios_heading = doc.add_paragraph()
                scenarios_heading.add_run("MODELOS TREINADOS POR CENÁRIO: ").bold = True

                best_scenario = None
                best_score = 0

                for scenario, scenario_data in mt_data['scenarios'].items():
                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                        results = scenario_data['predictor'].results

                        doc.add_paragraph('')
                        scenario_para = doc.add_paragraph()
                        scenario_para.add_run(f"CENÁRIO: {scenario}").bold = True

                        acc = results.get('balanced_accuracy', 0)
                        f1 = results.get('f1_score', 0)
                        auc = results.get('auc_score', 0)
                        threshold = results.get('optimal_threshold', 0.5)

                        doc.add_paragraph(f"• Acurácia Balanceada: {acc:.4f} ({acc:.1%})")
                        doc.add_paragraph(f"• F1-Score: {f1:.4f}")
                        doc.add_paragraph(f"• AUC: {auc:.4f}")
                        doc.add_paragraph(f"• Threshold Ótimo: {threshold:.3f}")

                        meta_model = scenario_data.get('best_meta_name', 'N/A')
                        features_used = scenario_data.get('features_used', [])

                        doc.add_paragraph(f"• Meta-modelo: {meta_model}")
                        doc.add_paragraph(f"• Features utilizadas: {len(features_used)}")

                        sensitivity, specificity = self._calculate_sensitivity_specificity(results)
                        if sensitivity > 0 or specificity > 0:
                            doc.add_paragraph("MÉTRICAS CLÍNICAS:")
                            doc.add_paragraph(f"• Sensibilidade: {sensitivity:.4f}")
                            doc.add_paragraph(f"• Especificidade: {specificity:.4f}")

                        if acc > best_score:
                            best_score = acc
                            best_scenario = scenario

                if best_scenario:
                    doc.add_paragraph('')
                    best_para = doc.add_paragraph()
                    best_para.add_run(f"🏆 MELHOR MODELO: {best_scenario}").bold = True
                    doc.add_paragraph(f"   Acurácia: {best_score:.1%}")

            elif 'results' in mt_data:
                results = mt_data['results']
                doc.add_paragraph("PERFORMANCE DO MODELO PRINCIPAL:")
                doc.add_paragraph(f"• Acurácia Balanceada: {results.get('balanced_accuracy', 0):.1%}")
                doc.add_paragraph(f"• F1-Score: {results.get('f1_score', 0):.4f}")
                doc.add_paragraph(f"• AUC: {results.get('auc_score', 0):.4f}")

            else:
                doc.add_paragraph("Estrutura de dados de performance não reconhecida.")

        except Exception as e:
            logger.warning(f"⚠️ Erro na performance: {e}")
            doc.add_paragraph("Erro ao processar dados de performance. Verifique logs técnicos.")

    def _calculate_sensitivity_specificity(self, results: Dict) -> Tuple[float, float]:
        """Calcula sensibilidade e especificidade"""

        try:
            cm_key = 'optimized_confusion_matrix'
            if cm_key in results:
                cm = results[cm_key]
                if hasattr(cm, 'ravel') and cm.size == 4:
                    tn, fp, fn, tp = cm.ravel()
                    sensitivity = tp / (tp + fn) if (tp + fn) > 0 else 0.0
                    specificity = tn / (tn + fp) if (tn + fp) > 0 else 0.0
                    return sensitivity, specificity
            return 0.0, 0.0
        except Exception:
            return 0.0, 0.0

    def _add_implementation_roadmap(self, doc):
        """Roadmap de implementação"""

        try:
            doc.add_heading('FEATURES UTILIZADAS', 1)

            if 'feature_selection' in self.results:
                fs_data = self.results['feature_selection']
                scenarios = fs_data.get('scenarios_processed', [])

                doc.add_paragraph("FEATURES SELECIONADAS POR CENÁRIO:")

                if scenarios and hasattr(fs_data.get('selector'), 'get_data_for_scenario'):
                    for scenario in scenarios:
                        try:
                            data, features = fs_data['selector'].get_data_for_scenario(scenario)

                            doc.add_paragraph('')
                            scenario_para = doc.add_paragraph()
                            scenario_para.add_run(f"{scenario} ({len(features)} features): ").bold = True

                            for i, feature in enumerate(features[:20], 1):
                                doc.add_paragraph(f"{i:2d}. {feature}")

                            if len(features) > 20:
                                doc.add_paragraph(f"... e mais {len(features) - 20} features")

                        except Exception as e:
                            doc.add_paragraph(f"{scenario}: Erro ao obter features")

            if 'model_training' in self.results:
                mt_data = self.results['model_training']

                if 'scenarios' in mt_data:
                    doc.add_paragraph('')
                    doc.add_paragraph("FEATURES EFETIVAMENTE UTILIZADAS NO TREINAMENTO:")

                    for scenario, scenario_data in mt_data['scenarios'].items():
                        features_used = scenario_data.get('features_used', [])
                        if features_used:
                            doc.add_paragraph('')
                            scenario_para = doc.add_paragraph()
                            scenario_para.add_run(f"{scenario}: {len(features_used)} features").bold = True

                            if 'predictor' in scenario_data:
                                predictor = scenario_data['predictor']
                                if hasattr(predictor, 'get_feature_importance'):
                                    try:
                                        importance = predictor.get_feature_importance()
                                        if importance:
                                            sorted_importance = sorted(
                                                importance.items(),
                                                key=lambda x: x[1],
                                                reverse=True
                                            )

                                            doc.add_paragraph("Top 20 features mais importantes:")
                                            for i, (feature, score) in enumerate(sorted_importance[:20], 1):
                                                doc.add_paragraph(f"{i:2d}. {feature}: {score:.4f}")
                                    except:
                                        pass

            doc.add_heading('RECOMENDAÇÕES DETALHADAS', 1)

            risk_level = self._get_overall_risk()
            recommendation = self._get_final_recommendation()

            doc.add_paragraph('')

            if risk_level == 'CRITICAL':
                doc.add_paragraph("🚨 AÇÃO IMEDIATA NECESSÁRIA:")
                doc.add_paragraph("• Revisar completamente o dataset")
                doc.add_paragraph("• Reprocessar dados antes de novo treinamento")

            elif risk_level == 'HIGH':
                doc.add_paragraph("⚠️ IMPLEMENTAÇÃO COM EXTREMA CAUTELA:")
                doc.add_paragraph("• Monitoramento rigoroso das predições")

            elif risk_level in ['MODERATE', 'LOW']:
                doc.add_paragraph("✅ IMPLEMENTAÇÃO RECOMENDADA:")
                doc.add_paragraph("• Sistema validado para uso clínico")

        except Exception as e:
            logger.warning(f"⚠️ Erro no roadmap: {e}")

    def _get_overall_status(self) -> str:
        """Status geral do sistema"""

        completed = len(self.results)
        total = 5

        if completed == total:
            return "TOTALMENTE CONCLUÍDO"
        elif completed >= 4:
            return "PRATICAMENTE CONCLUÍDO"
        elif completed >= 3:
            return "PARCIALMENTE CONCLUÍDO"
        elif completed >= 2:
            return "EM PROGRESSO"
        else:
            return "INICIADO"

    def _get_overall_risk(self) -> str:
        """Risco geral do sistema"""

        if 'leakage_detection' in self.results:
            ld_data = self.results['leakage_detection']
            return ld_data.get('overall_risk_level', 'UNKNOWN')
        return 'UNKNOWN'

    def _get_final_recommendation(self) -> str:
        """Recomendação final baseada em múltiplos fatores"""

        risk = self._get_overall_risk()
        best_accuracy = self._get_best_accuracy()
        completed_steps = len(self.results)

        risk_score = {
            'LOW': 4, 'MODERATE': 3, 'HIGH': 2, 'CRITICAL': 1, 'UNKNOWN': 2
        }.get(risk, 2)

        if best_accuracy >= 0.80:
            perf_score = 4
        elif best_accuracy >= 0.70:
            perf_score = 3
        elif best_accuracy >= 0.60:
            perf_score = 2
        else:
            perf_score = 1

        completeness_score = min(4, completed_steps)

        total_score = (risk_score + perf_score + completeness_score) / 3

        if total_score >= 3.5:
            return "IMPLEMENTAÇÃO COMPLETA RECOMENDADA"
        elif total_score >= 2.5:
            return "IMPLEMENTAÇÃO GRADUAL RECOMENDADA"
        elif total_score >= 2.0:
            return "PILOTO CONTROLADO RECOMENDADO"
        else:
            return "DESENVOLVIMENTO ADICIONAL NECESSÁRIO"

    def _get_best_accuracy(self) -> float:
        """Melhor acurácia obtida"""

        try:
            if 'model_training' not in self.results:
                return 0.0

            mt_data = self.results['model_training']
            best_acc = 0.0

            if 'scenarios' in mt_data:
                for scenario_data in mt_data['scenarios'].values():
                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                        acc = scenario_data['predictor'].results.get('balanced_accuracy', 0)
                        best_acc = max(best_acc, acc)
            elif 'results' in mt_data:
                best_acc = mt_data['results'].get('balanced_accuracy', 0)

            return best_acc
        except Exception:
            return 0.0

    def _get_leakage_interpretation(self, risk_level: str) -> str:
        """Interpretação detalhada do nível de risco"""

        interpretations = {
            'LOW': 'Dados validados com baixo risco de vazamento temporal.',
            'MODERATE': 'Risco moderado detectado. Performance pode estar ligeiramente inflada.',
            'HIGH': 'Alto risco de vazamento temporal identificado.',
            'CRITICAL': 'Vazamento crítico detectado que compromete completamente a validade do modelo.',
            'UNKNOWN': 'Análise de risco não completada ou dados insuficientes para determinação.'
        }
        return interpretations.get(risk_level, 'Interpretação não disponível para este nível de risco.')

    def _write_fallback_content(self, f):
        """Conteúdo básico para relatório fallback"""

        try:
            f.write(f"Data de Execução: {self.timestamp.strftime('%d/%m/%Y %H:%M:%S')}\n")
            f.write(f"Pipeline Executado: {len(self.results)} de 5 etapas\n")
            f.write(f"Status Geral: {self._get_overall_status()}\n\n")
            f.write("ETAPAS EXECUTADAS:\n")
            f.write("-" * 30 + "\n")

            step_names = {
                'feature_selection': 'Seleção de Features',
                'clinical_analysis': 'Análise Temporal',
                'leakage_detection': 'Detecção de Vazamento',
                'model_training': 'Treinamento do Modelo',
                'report_generation': 'Geração de Relatórios'
            }

            for step_key, step_name in step_names.items():
                if step_key in self.results:
                    f.write(f"✓ {step_name}: CONCLUÍDA\n")
                else:
                    f.write(f"✗ {step_name}: NÃO EXECUTADA\n")

            f.write(f"\nRESULTADOS PRINCIPAIS:\n")
            f.write("-" * 30 + "\n")
            f.write(f"Risco Detectado: {self._get_overall_risk()}\n")
            f.write(f"Melhor Acurácia: {self._get_best_accuracy():.1%}\n")
            f.write(f"Recomendação: {self._get_final_recommendation()}\n")

            if 'feature_selection' in self.results:
                fs_data = self.results['feature_selection']
                scenarios = fs_data.get('scenarios_processed', [])
                if scenarios:
                    f.write(f"\nCenários Processados: {len(scenarios)}\n")
                    for scenario in scenarios:
                        f.write(f"  - {scenario}\n")

            f.write(f"\nRelatório gerado automaticamente pelo Sistema IA Tuberculose Infantil\n")
            f.write(f"Desenvolvido por: Janduhy Finizola da Cunha Neto\n")

        except Exception as e:
            f.write(f"Erro ao gerar conteúdo fallback: {e}\n")

    def generate_technical_report(self) -> str:
        """Gera relatório técnico em texto"""

        try:
            logger.info("🔧 Gerando relatório técnico...")

            timestamp = self.timestamp.strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / "reports" / f"relatorio_tecnico_{timestamp}.txt"

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("="*80 + "\n")
                f.write("RELATÓRIO TÉCNICO - IA TUBERCULOSE INFANTIL\n")
                f.write(f"Gerado: {self.timestamp.strftime('%d/%m/%Y %H:%M:%S')}\n")
                f.write("="*80 + "\n\n")
                f.write("1. DIAGNÓSTICO DO SISTEMA\n")
                f.write("-"*40 + "\n")
                f.write(f"Etapas executadas: {len(self.results)}/5\n")
                f.write(f"Etapas disponíveis: {list(self.results.keys())}\n")
                f.write(f"Erros encontrados: {len(self.data_summary.get('errors_found', []))}\n")
                f.write(f"Status geral: {self._get_overall_status()}\n\n")

                self._write_feature_selection_technical(f)
                self._write_leakage_analysis_technical(f)
                self._write_model_performance_technical(f)
                self._write_recommendations_technical(f)

            logger.info(f"✅ Relatório técnico: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"❌ Erro relatório técnico: {e}")
            return ""

    def _write_feature_selection_technical(self, f):
        """Seção técnica de seleção de features"""

        f.write("2. SELEÇÃO DE FEATURES - ANÁLISE TÉCNICA\n")
        f.write("-"*40 + "\n")

        if 'feature_selection' not in self.results:
            f.write("Seleção de features não executada.\n\n")
            return

        fs_data = self.results['feature_selection']
        scenarios = fs_data.get('scenarios_processed', [])
        derived = fs_data.get('derived_features', [])

        f.write(f"Cenários processados: {len(scenarios)}\n")
        f.write(f"Features derivadas: {len(derived)}\n")
        f.write(f"Configuração utilizada: {fs_data.get('config_used', {})}\n\n")

        if scenarios:
            f.write("Detalhamento por cenário:\n")
            for scenario in scenarios:
                f.write(f"  {scenario}:\n")
                # Tentar obter detalhes se disponível
                if hasattr(fs_data.get('selector'), 'get_data_for_scenario'):
                    try:
                        data, features = fs_data['selector'].get_data_for_scenario(scenario)
                        f.write(f"    Registros: {len(data)}\n")
                        f.write(f"    Features selecionadas: {len(features)}\n")
                        f.write(f"    Primeiras 10 features: {features[:10]}\n")
                    except Exception as e:
                        f.write(f"    Erro ao obter detalhes: {e}\n")
                f.write("\n")

        f.write("\n")

    def _write_leakage_analysis_technical(self, f):
        """Seção técnica de análise de vazamento"""

        f.write("3. ANÁLISE DE VAZAMENTO TEMPORAL - DETALHES TÉCNICOS\n")
        f.write("-"*40 + "\n")

        if 'leakage_detection' not in self.results:
            f.write("Análise de vazamento não executada.\n\n")
            return

        ld_data = self.results['leakage_detection']
        overall_risk = ld_data.get('overall_risk_level', 'UNKNOWN')

        f.write(f"Risco geral detectado: {overall_risk}\n")
        f.write(f"Cenários analisados: {len(ld_data.get('scenarios_analyzed', []))}\n\n")

        if 'results_by_scenario' in ld_data:
            f.write("Resultados por cenário:\n")
            for scenario, results in ld_data['results_by_scenario'].items():
                f.write(f"  {scenario}:\n")
                f.write(f"    Risco: {results.get('risk_level', 'UNKNOWN')}\n")
                f.write(f"    Features suspeitas: {results.get('suspicious_count', 0)}\n")
                f.write(f"    Total features: {results.get('total_features', 0)}\n")

                if 'features_analysis' in results:
                    suspicious_features = [f for f, analysis in results['features_analysis'].items()
                                         if analysis.get('status') == 'SUSPEITA']
                    if suspicious_features:
                        f.write(f"    Features suspeitas específicas: {suspicious_features}\n")
                f.write("\n")

        f.write(f"Interpretação técnica:\n")
        f.write(f"  {self._get_leakage_interpretation(overall_risk)}\n\n")

    def _write_model_performance_technical(self, f):
        """Seção técnica de performance do modelo"""

        f.write("4. PERFORMANCE DOS MODELOS - ANÁLISE TÉCNICA\n")
        f.write("-"*40 + "\n")

        if 'model_training' not in self.results:
            f.write("Treinamento de modelos não executado.\n\n")
            return

        mt_data = self.results['model_training']

        config_used = mt_data.get('config_used', {})
        f.write("Configuração de treinamento:\n")
        for key, value in config_used.items():
            f.write(f"  {key}: {value}\n")
        f.write("\n")

        if 'scenarios' in mt_data:
            f.write(f"Modelos treinados: {len(mt_data['scenarios'])}\n\n")

            for scenario, scenario_data in mt_data['scenarios'].items():
                f.write(f"{scenario}:\n")

                if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                    results = scenario_data['predictor'].results

                    f.write(f"  Acurácia balanceada: {results.get('balanced_accuracy', 0):.4f}\n")
                    f.write(f"  F1-Score: {results.get('f1_score', 0):.4f}\n")
                    f.write(f"  AUC: {results.get('auc_score', 0):.4f}\n")
                    f.write(f"  Threshold ótimo: {results.get('optimal_threshold', 0.5):.4f}\n")

                    meta_model = scenario_data.get('best_meta_name', 'N/A')
                    features_used = scenario_data.get('features_used', [])
                    f.write(f"  Meta-modelo selecionado: {meta_model}\n")
                    f.write(f"  Features utilizadas: {len(features_used)}\n")

                    sensitivity, specificity = self._calculate_sensitivity_specificity(results)
                    if sensitivity > 0 or specificity > 0:
                        f.write(f"  Sensibilidade: {sensitivity:.4f}\n")
                        f.write(f"  Especificidade: {specificity:.4f}\n")

                else:
                    f.write("  Dados de performance não disponíveis\n")

                f.write("\n")

        best_scenario = mt_data.get('best_scenario')
        if best_scenario:
            f.write(f"Melhor modelo identificado: {best_scenario}\n")
            best_accuracy = self._get_best_accuracy()
            f.write(f"Melhor acurácia: {best_accuracy:.4f}\n\n")

    def _write_recommendations_technical(self, f):
        """Seção técnica de recomendações"""

        f.write("5. RECOMENDAÇÕES TÉCNICAS\n")
        f.write("-"*40 + "\n")

        recommendation = self._get_final_recommendation()
        risk = self._get_overall_risk()
        best_acc = self._get_best_accuracy()

        f.write(f"Recomendação principal: {recommendation}\n")
        f.write(f"Justificativa baseada em:\n")
        f.write(f"  - Nível de risco: {risk}\n")
        f.write(f"  - Melhor performance: {best_acc:.1%}\n")
        f.write(f"  - Etapas concluídas: {len(self.results)}/5\n\n")

        if 'model_training' in self.results:
            mt_data = self.results['model_training']
            if 'scenarios' in mt_data:
                f.write("Thresholds recomendados por cenário:\n")
                for scenario, scenario_data in mt_data['scenarios'].items():
                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                        threshold = scenario_data['predictor'].results.get('optimal_threshold', 0.5)
                        f.write(f"  {scenario}: {threshold:.4f}\n")
                f.write("\n")

    def generate_executive_summary(self) -> str:
        """Gera resumo executivo"""

        try:
            logger.info("📋 Gerando resumo executivo...")

            timestamp = self.timestamp.strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / "reports" / f"resumo_executivo_{timestamp}.txt"

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("="*80 + "\n")
                f.write("RESUMO EXECUTIVO - SISTEMA IA TUBERCULOSE INFANTIL\n")
                f.write(f"Data: {self.timestamp.strftime('%d/%m/%Y %H:%M:%S')}\n")
                f.write("="*80 + "\n\n")

                f.write("OBJETIVO:\n")
                f.write("Desenvolvimento de sistema de inteligência artificial para predição\n")
                f.write("de sobrevivência em casos de tuberculose infantil, com validação\n")
                f.write("rigorosa anti-vazamento temporal.\n\n")

                f.write("STATUS DA EXECUÇÃO:\n")
                f.write(f"Pipeline: {len(self.results)}/5 etapas concluídas\n")
                f.write(f"Status: {self._get_overall_status()}\n\n")

                f.write("PRINCIPAIS RESULTADOS:\n")

                if 'feature_selection' in self.results:
                    fs_data = self.results['feature_selection']
                    scenarios = fs_data.get('scenarios_processed', [])
                    f.write(f"• Cenários demográficos analisados: {len(scenarios)}\n")

                best_acc = self._get_best_accuracy()
                if best_acc > 0:
                    f.write(f"• Melhor performance alcançada: {best_acc:.1%}\n")

                risk = self._get_overall_risk()
                f.write(f"• Nível de risco de vazamento: {risk}\n")

                recommendation = self._get_final_recommendation()
                f.write(f"• Recomendação: {recommendation}\n\n")

                f.write("RECOMENDAÇÃO FINAL:\n")
                f.write(f"{recommendation}\n\n")

                f.write("Relatório preparado pelo Sistema de IA Tuberculose Infantil\n")
                f.write("Desenvolvido por: Janduhy Finizola da Cunha Neto\n")

            logger.info(f"✅ Resumo executivo: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"❌ Erro resumo executivo: {e}")
            return ""

    def generate_visualizations(self) -> List[str]:
        """Gera visualizações básicas"""

        try:
            logger.info("📊 Gerando visualizações...")

            viz_paths = []
            viz_dir = self.output_dir / "visualizations"
            viz_dir.mkdir(exist_ok=True)

            if 'model_training' in self.results:
                path = self._create_performance_chart(viz_dir)
                if path: viz_paths.append(path)

                path = self._create_confusion_matrix(viz_dir)
                if path: viz_paths.append(path)

            path = self._create_risk_chart(viz_dir)
            if path: viz_paths.append(path)

            path = self._create_pipeline_status(viz_dir)
            if path: viz_paths.append(path)

            logger.info(f"✅ {len(viz_paths)} visualizações geradas")
            return viz_paths

        except Exception as e:
            logger.error(f"❌ Erro visualizações: {e}")
            return []

    def _create_performance_chart(self, viz_dir: Path) -> Optional[str]:
        """Gráfico de performance"""

        try:
            if 'model_training' not in self.results:
                return None

            mt_data = self.results['model_training']

            scenarios_data = {}

            if 'scenarios' in mt_data:
                for scenario, scenario_data in mt_data['scenarios'].items():
                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                        results = scenario_data['predictor'].results
                        scenarios_data[scenario] = {
                            'accuracy': results.get('balanced_accuracy', 0),
                            'f1': results.get('f1_score', 0),
                            'auc': results.get('auc_score', 0)
                        }
            elif 'results' in mt_data:
                results = mt_data['results']
                scenarios_data['Modelo Principal'] = {
                    'accuracy': results.get('balanced_accuracy', 0),
                    'f1': results.get('f1_score', 0),
                    'auc': results.get('auc_score', 0)
                }

            if not scenarios_data:
                return None

            fig, axes = plt.subplots(1, len(scenarios_data), figsize=(15, 6))
            if len(scenarios_data) == 1:
                axes = [axes]

            for i, (scenario, metrics) in enumerate(scenarios_data.items()):
                ax = axes[i] if len(scenarios_data) > 1 else axes[0]

                metric_names = ['Accuracy', 'F1-Score', 'AUC']
                values = [metrics['accuracy'], metrics['f1'], metrics['auc']]

                bars = ax.bar(metric_names, values, color=['#3498db', '#e74c3c', '#2ecc71'])

                for bar, value in zip(bars, values):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01,
                           f'{value:.3f}', ha='center', va='bottom', fontweight='bold')

                ax.set_title(f'Performance - {scenario}', fontsize=14, fontweight='bold')
                ax.set_ylabel('Score')
                ax.set_ylim(0, 1.1)
                ax.grid(axis='y', alpha=0.3)

            plt.tight_layout()
            filepath = viz_dir / f"performance_{self.timestamp.strftime('%Y%m%d_%H%M%S')}.png"
            plt.savefig(filepath, dpi=300, bbox_inches='tight')
            plt.close()

            return str(filepath)

        except Exception as e:
            logger.warning(f"⚠️ Erro gráfico performance: {e}")
            return None

    def _create_confusion_matrix(self, viz_dir: Path) -> Optional[str]:
        """Matriz de confusão"""

        try:
            if 'model_training' not in self.results:
                return None

            mt_data = self.results['model_training']

            cm = None
            scenario_name = "Principal"

            if 'scenarios' in mt_data:
                best_scenario = mt_data.get('best_scenario')
                if best_scenario and best_scenario in mt_data['scenarios']:
                    scenario_data = mt_data['scenarios'][best_scenario]
                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                        cm = scenario_data['predictor'].results.get('optimized_confusion_matrix')
                        scenario_name = best_scenario

                if cm is None:
                    for scenario, scenario_data in mt_data['scenarios'].items():
                        if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                            cm = scenario_data['predictor'].results.get('optimized_confusion_matrix')
                            scenario_name = scenario
                            break

            elif 'results' in mt_data:
                cm = mt_data['results'].get('optimized_confusion_matrix')

            if cm is None:
                return None

            plt.figure(figsize=(8, 6))
            sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
                       xticklabels=['Óbito', 'Sobreviveu'],
                       yticklabels=['Óbito', 'Sobreviveu'])

            plt.title(f'Matriz de Confusão - {scenario_name}', fontsize=14, fontweight='bold')
            plt.xlabel('Predição')
            plt.ylabel('Real')

            filepath = viz_dir / f"confusion_{self.timestamp.strftime('%Y%m%d_%H%M%S')}.png"
            plt.savefig(filepath, dpi=300, bbox_inches='tight')
            plt.close()

            return str(filepath)

        except Exception as e:
            logger.warning(f"⚠️ Erro matriz confusão: {e}")
            return None

    def _create_risk_chart(self, viz_dir: Path) -> Optional[str]:
        """Gráfico de risco"""

        try:
            risks = {}

            if 'leakage_detection' in self.results:
                ld_data = self.results['leakage_detection']
                risks['Geral'] = ld_data.get('overall_risk_level', 'UNKNOWN')

                if 'results_by_scenario' in ld_data:
                    for scenario, results in ld_data['results_by_scenario'].items():
                        risks[scenario] = results.get('risk_level', 'UNKNOWN')

            if not risks:
                return None

            risk_mapping = {'LOW': 1, 'MODERATE': 2, 'HIGH': 3, 'CRITICAL': 4, 'UNKNOWN': 0}
            colors = ['gray', '#2ecc71', '#f39c12', '#e74c3c', '#8e44ad']

            stages = list(risks.keys())
            risk_values = [risk_mapping.get(risks[stage], 0) for stage in stages]
            bar_colors = [colors[val] for val in risk_values]

            plt.figure(figsize=(12, 6))
            bars = plt.bar(stages, risk_values, color=bar_colors, alpha=0.8, edgecolor='black', linewidth=1)

            for bar, stage in zip(bars, stages):
                risk_level = risks[stage]
                plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                         risk_level, ha='center', va='bottom', fontweight='bold', fontsize=10)

            plt.title('Análise de Risco de Data Leakage por Cenário', fontsize=16, fontweight='bold', pad=20)
            plt.ylabel('Nível de Risco', fontsize=12, fontweight='bold')
            plt.xlabel('Cenários', fontsize=12, fontweight='bold')
            plt.yticks([1, 2, 3, 4], ['LOW', 'MODERATE', 'HIGH', 'CRITICAL'])
            plt.ylim(0, 5)
            plt.xticks(rotation=45, ha='right')
            plt.grid(axis='y', alpha=0.3, linestyle='--')
            plt.tight_layout()

            risk_labels = ['LOW', 'MODERATE', 'HIGH', 'CRITICAL']
            risk_colors = ['#2ecc71', '#f39c12', '#e74c3c', '#8e44ad']
            legend_elements = [plt.Rectangle((0, 0), 1, 1, facecolor=color, alpha=0.8)
                               for color in risk_colors]
            plt.legend(legend_elements, risk_labels, loc='upper right', title='Níveis de Risco')

            filepath = viz_dir / f"risk_analysis_{self.timestamp.strftime('%Y%m%d_%H%M%S')}.png"
            plt.savefig(filepath, dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()

            return str(filepath)

        except Exception as e:
            logger.warning(f"⚠️ Erro gráfico risco: {e}")
            return None

    def _create_pipeline_status(self, viz_dir: Path) -> Optional[str]:
        """Status do pipeline com design profissional"""

        try:
            stages = ['Seleção\nFeatures', 'Análise\nTemporal', 'Detecção\nLeakage', 'Treinamento\nModelo',
                      'Geração\nRelatórios']
            keys = ['feature_selection', 'clinical_analysis', 'leakage_detection', 'model_training',
                    'report_generation']

            status = [1 if key in self.results else 0 for key in keys]
            colors = ['#2ecc71' if s == 1 else '#e74c3c' for s in status]

            fig, ax = plt.subplots(figsize=(14, 8))

            y_pos = range(len(stages))
            bars = ax.barh(y_pos, [1] * len(stages), color=['#ecf0f1'] * len(stages), alpha=0.3, height=0.6)

            progress_bars = ax.barh(y_pos, status, color=colors, alpha=0.8, height=0.6,
                                    edgecolor='black', linewidth=1)

            for i, (bar, s, stage) in enumerate(zip(progress_bars, status, stages)):
                symbol = '✅' if s == 1 else '❌'
                status_text = 'CONCLUÍDA' if s == 1 else 'PENDENTE'

                ax.text(0.05, i, symbol, fontsize=20, va='center', ha='left')

                ax.text(0.15, i, stage.replace('\n', ' '), fontsize=12, fontweight='bold',
                        va='center', ha='left')

                ax.text(0.85, i, status_text, fontsize=10, fontweight='bold',
                        va='center', ha='right', color='white' if s == 1 else 'black')

            ax.set_xlim(0, 1)
            ax.set_ylim(-0.5, len(stages) - 0.5)
            ax.set_yticks([])
            ax.set_xticks([])

            for spine in ax.spines.values():
                spine.set_visible(False)

            completed_count = sum(status)
            ax.set_title(f'Status do Pipeline de IA - {completed_count}/{len(stages)} Etapas Concluídas',
                         fontsize=16, fontweight='bold', pad=30)

            progress_percentage = (completed_count / len(stages)) * 100
            fig.text(0.5, 0.02, f'Progresso Geral: {progress_percentage:.1f}% | Sistema: IA Tuberculose Infantil',
                     ha='center', fontsize=10, style='italic')

            plt.tight_layout()
            filepath = viz_dir / f"pipeline_status_{self.timestamp.strftime('%Y%m%d_%H%M%S')}.png"
            plt.savefig(filepath, dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()

            return str(filepath)

        except Exception as e:
            logger.warning(f"⚠️ Erro status pipeline: {e}")
            return None

    def export_summary_json(self) -> str:
        """Exporta resumo em JSON"""

        def safe_convert(obj):
            """Converte objetos para JSON"""
            if isinstance(obj, np.ndarray):
                return obj.tolist()
            elif isinstance(obj, (np.integer, np.int64, np.int32)):
                return int(obj)
            elif isinstance(obj, (np.floating, np.float64, np.float32)):
                return float(obj)
            elif isinstance(obj, pd.DataFrame):
                return obj.to_dict('records')
            elif hasattr(obj, 'isoformat'):
                return obj.isoformat()
            elif hasattr(obj, '__dict__'):
                return str(obj)
            else:
                return str(obj)

        def safe_serialize(obj):
            """Serializa recursivamente com tratamento de erro"""

            try:
                if obj is None:
                    return None
                elif isinstance(obj, dict):
                    return {str(k): safe_serialize(v) for k, v in obj.items()}
                elif isinstance(obj, (list, tuple)):
                    return [safe_serialize(item) for item in obj]
                elif isinstance(obj, (str, int, float, bool)):
                    return obj
                else:
                    return safe_convert(obj)
            except Exception as e:
                logger.warning(f"Erro ao serializar objeto: {e}")
                return str(obj)

        try:
            summary = {
                'metadata': {
                    'timestamp': self.timestamp.isoformat(),
                    'system_name': 'Sistema IA Tuberculose Infantil',
                    'version': '1.0.0',
                    'author': 'Janduhy Finizola da Cunha Neto',
                    'report_type': 'Pipeline Execution Summary'
                },
                'execution_status': {
                    'completed_steps': len(self.results),
                    'total_steps': 5,
                    'completion_rate': len(self.results) / 5 * 100,
                    'steps_executed': list(self.results.keys()),
                    'overall_status': self._get_overall_status()
                },
                'risk_assessment': {
                    'overall_risk_level': self._get_overall_risk(),
                    'risk_interpretation': self._get_leakage_interpretation(self._get_overall_risk()),
                    'validation_status': 'PASSED' if self._get_overall_risk() in ['LOW', 'MODERATE'] else 'FAILED'
                },
                'model_performance': self._get_comprehensive_performance_data(),
                'recommendations': {
                    'final_recommendation': self._get_final_recommendation(),
                    'implementation_ready': self._get_overall_risk() in ['LOW', 'MODERATE'],
                    'next_steps': self._generate_next_steps()
                },
                'scenarios_analysis': self._get_detailed_scenarios_summary(),
                'feature_analysis': self._get_comprehensive_feature_summary(),
                'quality_metrics': self._calculate_quality_metrics()
            }

            summary_safe = safe_serialize(summary)

            timestamp = self.timestamp.strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / "reports" / f"executive_summary_{timestamp}.json"

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(summary_safe, f, indent=2, ensure_ascii=False)

            logger.info(f"✅ JSON executivo exportado: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"❌ Erro ao exportar JSON: {e}")

            error_summary = {
                'status': 'ERROR',
                'error_type': type(e).__name__,
                'error_message': str(e),
                'timestamp': self.timestamp.isoformat(),
                'completed_steps': len(self.results),
                'partial_data': {
                    'risk_level': self._get_overall_risk(),
                    'recommendation': self._get_final_recommendation(),
                    'steps_completed': list(self.results.keys())
                }
            }

            timestamp = self.timestamp.strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / "reports" / f"summary_error_{timestamp}.json"

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(error_summary, f, indent=2, ensure_ascii=False)

            return str(filepath)

    def _get_comprehensive_performance_data(self) -> Dict[str, Any]:
        """Dados abrangentes de performance"""

        performance_data = {
            'models_trained': 0,
            'best_accuracy': 0.0,
            'best_scenario': None,
            'scenarios_performance': {},
            'average_performance': {
                'accuracy': 0.0,
                'f1_score': 0.0,
                'auc_score': 0.0
            }
        }

        try:
            if 'model_training' not in self.results:
                return performance_data

            mt_data = self.results['model_training']

            if 'scenarios' in mt_data:
                performance_data['models_trained'] = len(mt_data['scenarios'])

                accuracies = []
                f1_scores = []
                auc_scores = []

                for scenario, scenario_data in mt_data['scenarios'].items():
                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                        results = scenario_data['predictor'].results

                        acc = results.get('balanced_accuracy', 0)
                        f1 = results.get('f1_score', 0)
                        auc = results.get('auc_score', 0)

                        performance_data['scenarios_performance'][scenario] = {
                            'accuracy': float(acc),
                            'f1_score': float(f1),
                            'auc_score': float(auc),
                            'threshold': float(results.get('optimal_threshold', 0.5)),
                            'meta_model': scenario_data.get('best_meta_name', 'N/A'),
                            'features_count': len(scenario_data.get('features_used', []))
                        }

                        accuracies.append(acc)
                        f1_scores.append(f1)
                        auc_scores.append(auc)

                        if acc > performance_data['best_accuracy']:
                            performance_data['best_accuracy'] = float(acc)
                            performance_data['best_scenario'] = scenario

                if accuracies:
                    performance_data['average_performance'] = {
                        'accuracy': float(np.mean(accuracies)),
                        'f1_score': float(np.mean(f1_scores)),
                        'auc_score': float(np.mean(auc_scores))
                    }

            elif 'results' in mt_data:
                performance_data['models_trained'] = 1
                results = mt_data['results']
                performance_data['best_accuracy'] = float(results.get('balanced_accuracy', 0))
                performance_data['best_scenario'] = 'Principal'

                performance_data['scenarios_performance']['Principal'] = {
                    'accuracy': float(results.get('balanced_accuracy', 0)),
                    'f1_score': float(results.get('f1_score', 0)),
                    'auc_score': float(results.get('auc_score', 0)),
                    'threshold': float(results.get('optimal_threshold', 0.5))
                }

        except Exception as e:
            logger.warning(f"Erro ao extrair dados de performance: {e}")

        return performance_data

    def _generate_next_steps(self) -> List[str]:
        """Gera próximos passos baseados no status atual"""

        next_steps = []

        risk = self._get_overall_risk()
        completed_steps = len(self.results)

        if completed_steps < 5:
            next_steps.append("Completar todas as etapas do pipeline")

        if risk == 'CRITICAL':
            next_steps.extend([
                "Revisar dataset para remover vazamento crítico",
                "Reprocessar seleção de features",
                "Executar nova rodada de treinamento"
            ])
        elif risk == 'HIGH':
            next_steps.extend([
                "Validar modelo em dataset externo",
                "Implementar monitoramento rigoroso",
                "Considerar retreinamento com features revisadas"
            ])
        elif risk in ['LOW', 'MODERATE']:
            next_steps.extend([
                "Preparar ambiente de produção",
                "Definir protocolos de monitoramento",
                "Treinar equipe para uso do sistema",
                "Implementar em piloto controlado"
            ])

        if 'model_training' in self.results:
            best_acc = self._get_best_accuracy()
            if best_acc < 0.70:
                next_steps.append("Considerar melhorias no modelo (mais dados, features adicionais)")
            elif best_acc >= 0.80:
                next_steps.append("Sistema pronto para implementação clínica")

        return next_steps

    def _get_detailed_scenarios_summary(self) -> Dict[str, Any]:
        """Resumo detalhado dos cenários"""

        summary = {
            'total_scenarios': 0,
            'scenarios_list': [],
            'scenarios_details': {},
            'derived_features_created': []
        }

        try:
            if 'feature_selection' in self.results:
                fs_data = self.results['feature_selection']
                scenarios = fs_data.get('scenarios_processed', [])

                summary['total_scenarios'] = len(scenarios)
                summary['scenarios_list'] = scenarios
                summary['derived_features_created'] = fs_data.get('derived_features', [])

                if scenarios and hasattr(fs_data.get('selector'), 'get_data_for_scenario'):
                    for scenario in scenarios:
                        try:
                            data, features = fs_data['selector'].get_data_for_scenario(scenario)

                            scenario_detail = {
                                'records_count': len(data),
                                'features_count': len(features),
                                'features_list': features
                            }

                            if 'SITUA_ENCE' in data.columns:
                                target_dist = data['SITUA_ENCE'].value_counts().to_dict()
                                scenario_detail['target_distribution'] = {
                                    str(k): int(v) for k, v in target_dist.items()
                                }

                            if 'model_training' in self.results:
                                mt_data = self.results['model_training']
                                if 'scenarios' in mt_data and scenario in mt_data['scenarios']:
                                    scenario_data = mt_data['scenarios'][scenario]
                                    if 'predictor' in scenario_data and hasattr(scenario_data['predictor'], 'results'):
                                        results = scenario_data['predictor'].results
                                        scenario_detail['performance'] = {
                                            'accuracy': float(results.get('balanced_accuracy', 0)),
                                            'f1_score': float(results.get('f1_score', 0)),
                                            'auc_score': float(results.get('auc_score', 0))
                                        }

                            summary['scenarios_details'][scenario] = scenario_detail

                        except Exception as e:
                            summary['scenarios_details'][scenario] = {'error': str(e)}

        except Exception as e:
            logger.warning(f"Erro ao gerar resumo de cenários: {e}")

        return summary

    def _get_comprehensive_feature_summary(self) -> Dict[str, Any]:
        """Resumo abrangente das features"""

        summary = {
            'derived_features': [],
            'derived_features_count': 0,
            'selection_config': {},
            'temporal_analysis_available': False,
            'leakage_analysis_available': False
        }

        try:
            if 'feature_selection' in self.results:
                fs_data = self.results['feature_selection']
                summary['derived_features'] = fs_data.get('derived_features', [])
                summary['derived_features_count'] = len(fs_data.get('derived_features', []))
                summary['selection_config'] = fs_data.get('config_used', {})

            if 'clinical_analysis' in self.results:
                summary['temporal_analysis_available'] = True
                ca_data = self.results['clinical_analysis']
                summary['temporal_features_safe'] = len(ca_data.get('safe_features', []))
                summary['temporal_features_risky'] = len(ca_data.get('temporal_features', []))

            if 'leakage_detection' in self.results:
                summary['leakage_analysis_available'] = True
                ld_data = self.results['leakage_detection']
                summary['overall_leakage_risk'] = ld_data.get('overall_risk_level', 'UNKNOWN')

        except Exception as e:
            logger.warning(f"Erro ao gerar resumo de features: {e}")

        return summary

    def _calculate_quality_metrics(self) -> Dict[str, Any]:
        """Calcula métricas de qualidade do pipeline"""

        quality = {
            'completeness_score': 0.0,
            'data_quality_score': 0.0,
            'model_quality_score': 0.0,
            'overall_quality_score': 0.0,
            'quality_level': 'UNKNOWN'
        }

        try:
            quality['completeness_score'] = len(self.results) / 5

            risk = self._get_overall_risk()
            risk_scores = {'LOW': 1.0, 'MODERATE': 0.8, 'HIGH': 0.4, 'CRITICAL': 0.0, 'UNKNOWN': 0.5}
            quality['data_quality_score'] = risk_scores.get(risk, 0.5)

            best_acc = self._get_best_accuracy()
            if best_acc >= 0.80:
                quality['model_quality_score'] = 1.0
            elif best_acc >= 0.70:
                quality['model_quality_score'] = 0.8
            elif best_acc >= 0.60:
                quality['model_quality_score'] = 0.6
            else:
                quality['model_quality_score'] = 0.4

            weights = [0.3, 0.4, 0.3]
            scores = [quality['completeness_score'], quality['data_quality_score'], quality['model_quality_score']]
            quality['overall_quality_score'] = sum(w * s for w, s in zip(weights, scores))

            overall = quality['overall_quality_score']
            if overall >= 0.8:
                quality['quality_level'] = 'EXCELENTE'
            elif overall >= 0.6:
                quality['quality_level'] = 'BOA'
            elif overall >= 0.4:
                quality['quality_level'] = 'MODERADA'
            else:
                quality['quality_level'] = 'BAIXA'

        except Exception as e:
            logger.warning(f"Erro ao calcular métricas de qualidade: {e}")

        return quality

def generate_comprehensive_reports(results: Dict[str, Any], output_dir: Path,
                                   config: Dict[str, Any]) -> Dict[str, Any]:
    """
    Função principal para gerar todos os relatórios

    Args:
        results: Resultados do pipeline de IA
        output_dir: Diretório de saída
        config: Configurações do sistema
    """

    try:
        generator = ComprehensiveReportGenerator(results, output_dir, config)

        generated_files = {}

        word_path = generator.generate_word_report()
        if word_path:
            generated_files['word_report'] = word_path

        tech_path = generator.generate_technical_report()
        if tech_path:
            generated_files['technical_report'] = tech_path

        exec_path = generator.generate_executive_summary()
        if exec_path:
            generated_files['executive_summary'] = exec_path

        json_path = generator.export_summary_json()
        if json_path:
            generated_files['json_export'] = json_path

        viz_paths = generator.generate_visualizations()
        if viz_paths:
            generated_files['visualizations'] = viz_paths

        return {
            'status': 'SUCCESS',
            'generated_files': generated_files,
            'total_reports': len(generated_files),
            'timestamp': generator.timestamp.isoformat(),
            'recommendation': generator._get_final_recommendation(),
            'risk_level': generator._get_overall_risk(),
            'best_accuracy': generator._get_best_accuracy()
        }

    except Exception as e:
        logger.error(f"❌ Erro na geração completa de relatórios: {e}")
        return {
            'status': 'ERROR',
            'error': str(e),
            'generated_files': {},
            'timestamp': datetime.now().isoformat()
        }


def create_emergency_report(results: Dict[str, Any], output_dir: Path,
                            error_info: str = None) -> str:
    """
    Cria relatório de emergência quando o sistema principal falha

    Args:
        results: Resultados parciais disponíveis
        output_dir: Diretório de saída
        error_info: Informações do erro ocorrido
    """

    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        emergency_path = Path(output_dir) / "reports" / f"relatorio_emergencia_{timestamp}.txt"

        with open(emergency_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("RELATÓRIO DE EMERGÊNCIA - SISTEMA IA TUBERCULOSE\n")
            f.write("=" * 80 + "\n")
            f.write(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
            if error_info:
                f.write(f"Motivo: {error_info}\n")
            f.write("=" * 80 + "\n\n")

            f.write("STATUS DO SISTEMA:\n")
            f.write("-" * 40 + "\n")
            f.write(f"Etapas disponíveis: {list(results.keys()) if results else 'Nenhuma'}\n")
            f.write(f"Total de etapas: {len(results) if results else 0}/5\n\n")

            if results:
                f.write("DADOS PARCIAIS DISPONÍVEIS:\n")
                f.write("-" * 40 + "\n")
                for step, data in results.items():
                    f.write(f"✓ {step}: {type(data).__name__}\n")

            f.write("\nRECOMENDAÇÃO:\n")
            f.write("-" * 40 + "\n")
            f.write("1. Verificar logs detalhados para identificar a causa\n")
            f.write("2. Reexecutar o pipeline a partir da última etapa válida\n")
            f.write("3. Contactar suporte técnico se o problema persistir\n\n")

            f.write("CONTATO DE SUPORTE:\n")
            f.write("-" * 40 + "\n")
            f.write("Desenvolvedor: Janduhy Finizola da Cunha Neto\n")
            f.write("Sistema: IA Tuberculose Infantil v1.0.0\n")

        return str(emergency_path)

    except Exception as e:
        logger.error(f"❌ Erro crítico ao gerar relatório de emergência: {e}")
        return ""

matplotlib.use('Agg')

pd.set_option('display.max_columns', None)
pd.set_option('display.width', None)
pd.set_option('display.max_colwidth', 50)

logger.info("✅ Módulo de relatórios profissionais carregado com sucesso")
logger.info("📊 Funcionalidades disponíveis:")
logger.info("   - Relatórios Word profissionais")
logger.info("   - Relatórios técnicos detalhados")
logger.info("   - Resumos executivos")
logger.info("   - Exports JSON estruturados")
logger.info("   - Visualizações avançadas")
logger.info("   - Relatórios de emergência")
logger.info("   - Fallbacks robustos")